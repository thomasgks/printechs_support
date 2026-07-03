# Copyright (c) 2026, Printechs and contributors
# License: MIT. See license.txt

"""Extract PDF/DOCX documents and generate PRAI FAQ records."""

from __future__ import annotations

import json
import re
from pathlib import Path

import frappe
from frappe import _
from frappe.utils import cint, get_files_path, sanitize_html

from printechs_support.printechs_support_system.api.prai_openai import is_openai_configured
from printechs_support.printechs_support_system.doctype.printechs_support_settings.printechs_support_settings import (
	get_prai_openai_config,
)

ALLOWED_EXTENSIONS = frozenset({".pdf", ".docx"})
MAX_EXTRACT_CHARS = 120_000
MAX_OPENAI_CHARS = 24_000
FAQ_CATEGORIES = frozenset(
	{
		"Modern POS",
		"ERPNext",
		"Promotions",
		"Stock",
		"Loyalty",
		"E-Wallet",
		"Troubleshooting",
		"General",
	}
)
MODULE_AREAS = frozenset(
	{
		"Modern POS",
		"Support",
		"WMS",
		"ERPNext",
		"Sales",
		"Purchase",
		"Stock",
		"Accounts",
		"HR",
		"General",
	}
)


def _resolve_file_path(file_url: str) -> str:
	path = (file_url or "").strip()
	if not path:
		frappe.throw(_("Attach a PDF or DOCX file first."), frappe.ValidationError)
	if path.startswith("/private/files/"):
		return get_files_path(*path.replace("/private/files/", "").split("/"), is_private=1)
	if path.startswith("/files/"):
		return get_files_path(*path.replace("/files/", "").split("/"), is_private=0)
	if path.startswith("http://") or path.startswith("https://"):
		frappe.throw(_("Remote file URLs are not supported. Upload the file to this site."), frappe.ValidationError)
	return path


def _extension(file_url: str) -> str:
	name = Path((file_url or "").split("?")[0]).name.lower()
	return Path(name).suffix


def extract_document_text(file_url: str) -> str:
	ext = _extension(file_url)
	if ext not in ALLOWED_EXTENSIONS:
		frappe.throw(
			_("Unsupported file type {0}. Upload PDF (.pdf) or Word (.docx) only.").format(ext or "(unknown)"),
			frappe.ValidationError,
		)
	file_path = _resolve_file_path(file_url)
	if ext == ".pdf":
		text = _extract_pdf_text(file_path)
	else:
		text = _extract_docx_text(file_path)
	text = re.sub(r"\n{3,}", "\n\n", text).strip()
	if not text:
		frappe.throw(_("No readable text was found in the document."), frappe.ValidationError)
	return text[:MAX_EXTRACT_CHARS]


def _extract_pdf_text(file_path: str) -> str:
	try:
		from pypdf import PdfReader
	except ImportError:
		frappe.throw(_("PDF support requires the pypdf package on the bench."), frappe.ValidationError)

	reader = PdfReader(file_path)
	parts: list[str] = []
	for page in reader.pages:
		parts.append(page.extract_text() or "")
	return "\n\n".join(part.strip() for part in parts if part and part.strip())


def _extract_docx_text(file_path: str) -> str:
	try:
		from docx import Document
	except ImportError:
		frappe.throw(_("DOCX support requires the python-docx package on the bench."), frappe.ValidationError)

	doc = Document(file_path)
	parts: list[str] = []
	for paragraph in doc.paragraphs:
		text = (paragraph.text or "").strip()
		if text:
			parts.append(text)
	for table in doc.tables:
		for row in table.rows:
			cells = [((cell.text or "").strip()) for cell in row.cells]
			cells = [cell for cell in cells if cell]
			if cells:
				parts.append(" | ".join(cells))
	return "\n\n".join(parts)


def _normalize_category(value: str, default: str) -> str:
	value = (value or "").strip()
	return value if value in FAQ_CATEGORIES else default


def _normalize_module_area(value: str, default: str) -> str:
	value = (value or "").strip()
	return value if value in MODULE_AREAS else default


def _normalize_faq_item(raw: dict, *, default_category: str, default_module_area: str) -> dict | None:
	title = re.sub(r"\s+", " ", str(raw.get("title") or "")).strip()
	question = re.sub(r"\s+", " ", str(raw.get("question") or "")).strip()
	keywords = re.sub(r"\s+", " ", str(raw.get("keywords") or "")).strip()
	answer = sanitize_html(str(raw.get("answer") or "")).strip()
	if not title or not answer:
		return None
	if not question:
		question = title
	if not keywords:
		keywords = ", ".join(_keyword_candidates(title, question))
	return {
		"title": title[:140],
		"question": question[:500],
		"keywords": keywords[:500],
		"answer": answer,
		"category": _normalize_category(str(raw.get("category") or ""), default_category),
		"module_area": _normalize_module_area(str(raw.get("module_area") or ""), default_module_area),
	}


def _keyword_candidates(*parts: str) -> list[str]:
	terms: list[str] = []
	seen: set[str] = set()
	for part in parts:
		for token in re.split(r"[^\w]+", (part or "").lower()):
			if len(token) < 3 or token in seen:
				continue
			seen.add(token)
			terms.append(token)
	return terms[:12]


def _parse_json_faqs(raw: str) -> list[dict]:
	text = (raw or "").strip()
	if not text:
		return []
	if text.startswith("```"):
		text = re.sub(r"^```(?:json)?\s*", "", text)
		text = re.sub(r"\s*```$", "", text)
	try:
		payload = json.loads(text)
	except json.JSONDecodeError:
		match = re.search(r"\{[\s\S]*\}", text)
		if not match:
			return []
		payload = json.loads(match.group(0))
	if isinstance(payload, list):
		return payload
	if isinstance(payload, dict):
		for key in ("faqs", "items", "questions", "data"):
			value = payload.get(key)
			if isinstance(value, list):
				return value
	return []


def _generate_faqs_with_openai(
	text: str,
	*,
	default_category: str,
	default_module_area: str,
) -> list[dict]:
	if not is_openai_configured():
		frappe.throw(
			_(
				"OpenAI is not configured. Enable OpenAI and add an API key in Printechs Support Settings, "
				"or turn off 'Use OpenAI to Generate FAQs' for basic section splitting."
			),
			frappe.ValidationError,
		)

	from printechs_support.printechs_support_system.api.prai_openai import _call_openai_chat

	cfg = get_prai_openai_config()
	api_key = cfg.get("api_key")
	model = cfg.get("model") or "gpt-4o-mini"
	if not api_key:
		frappe.throw(_("OpenAI API key is missing."), frappe.ValidationError)

	doc_text = text[:MAX_OPENAI_CHARS]
	categories = ", ".join(sorted(FAQ_CATEGORIES))
	module_areas = ", ".join(sorted(MODULE_AREAS))
	system_prompt = (
		"You convert Printechs support documentation into PRAI FAQ records for Modern POS, ERPNext, "
		"retail operations, and the Support Portal.\n"
		"Return ONLY valid JSON with this shape:\n"
		'{"faqs":[{"title":"...","question":"...","keywords":"comma,separated,terms",'
		'"answer":"<p>...</p><ol><li><strong>Step</strong> — detail</li></ol>",'
		'"category":"...", "module_area":"..."}]}\n'
		f"Allowed categories: {categories}\n"
		f"Allowed module areas: {module_areas}\n"
		"Rules:\n"
		"- Create as many useful FAQs as the document supports (procedures, troubleshooting, definitions).\n"
		"- Use HTML in answers with numbered steps where appropriate.\n"
		"- Keywords must help users find the FAQ in search.\n"
		"- Do not invent features not mentioned in the document.\n"
		f"- Default category: {default_category}. Default module area: {default_module_area}."
	)
	user_prompt = f"Document text:\n\n{doc_text}"
	raw = _call_openai_chat(
		messages=[
			{"role": "system", "content": system_prompt},
			{"role": "user", "content": user_prompt},
		],
		model=str(model),
		api_key=str(api_key),
		max_tokens=4096,
	)
	items = []
	for row in _parse_json_faqs(raw):
		if not isinstance(row, dict):
			continue
		normalized = _normalize_faq_item(
			row,
			default_category=default_category,
			default_module_area=default_module_area,
		)
		if normalized:
			items.append(normalized)
	if not items:
		frappe.throw(_("OpenAI did not return any usable FAQ entries."), frappe.ValidationError)
	return items


def _generate_faqs_basic(
	text: str,
	*,
	default_category: str,
	default_module_area: str,
) -> list[dict]:
	"""Fallback when OpenAI is disabled: split document into section-based FAQs."""
	chunks = re.split(r"\n(?=\d+[\.)]\s+)", text)
	if len(chunks) <= 1:
		chunks = re.split(r"\n(?=\#{1,3}\s+)", text)
	sections = [chunk.strip() for chunk in chunks if chunk.strip()]
	items: list[dict] = []
	for section in sections:
		lines = [line.strip() for line in section.splitlines() if line.strip()]
		if not lines:
			continue
		title = re.sub(r"^\d+[\.)]\s*", "", lines[0])
		title = re.sub(r"^#+\s*", "", title).strip()
		if len(title) < 8:
			continue
		body_lines = lines[1:] or lines
		answer_parts = []
		for line in body_lines[:12]:
			clean = re.sub(r"^\d+[\.)]\s*", "", line)
			clean = re.sub(r"^[-•]\s*", "", clean).strip()
			if clean:
				answer_parts.append(f"<li>{sanitize_html(clean)}</li>")
		if not answer_parts:
			continue
		answer = f"<p><strong>{sanitize_html(title)}</strong></p><ol>{''.join(answer_parts)}</ol>"
		normalized = _normalize_faq_item(
			{
				"title": title[:140],
				"question": title,
				"keywords": ", ".join(_keyword_candidates(title)),
				"answer": answer,
				"category": default_category,
				"module_area": default_module_area,
			},
			default_category=default_category,
			default_module_area=default_module_area,
		)
		if normalized:
			items.append(normalized)
	if not items:
		frappe.throw(
			_(
				"Could not detect FAQ sections automatically. Enable OpenAI generation for better results."
			),
			frappe.ValidationError,
		)
	return items[:25]


def generate_faq_items_from_text(
	text: str,
	*,
	default_category: str = "General",
	default_module_area: str = "General",
	use_openai: bool = True,
) -> list[dict]:
	if use_openai:
		return _generate_faqs_with_openai(
			text,
			default_category=default_category,
			default_module_area=default_module_area,
		)
	return _generate_faqs_basic(
		text,
		default_category=default_category,
		default_module_area=default_module_area,
	)


def preview_import_lines(items: list[dict]) -> list[dict]:
	lines: list[dict] = []
	for item in items:
		title = item["title"]
		existing = frappe.db.get_value("PRAI FAQ", {"title": title}, "name")
		action = "Update" if existing else "Create"
		lines.append(
			{
				"include": 1,
				"action": action,
				"title": title,
				"question": item.get("question") or title,
				"keywords": item.get("keywords") or "",
				"answer": item.get("answer") or "",
				"category": item.get("category") or "General",
				"module_area": item.get("module_area") or "General",
				"prai_faq": existing or "",
			}
		)
	return lines


def upsert_prai_faqs(items: list[dict], *, update_existing: bool = True) -> dict:
	created = 0
	updated = 0
	skipped = 0
	for item in items:
		title = item["title"]
		existing = frappe.db.get_value("PRAI FAQ", {"title": title}, "name")
		payload = {
			"question": item.get("question") or title,
			"keywords": item.get("keywords") or "",
			"answer": item.get("answer") or "",
			"category": item.get("category") or "General",
			"module_area": item.get("module_area") or "General",
			"is_active": 1,
		}
		if existing:
			if not update_existing:
				skipped += 1
				continue
			doc = frappe.get_doc("PRAI FAQ", existing)
			for key, value in payload.items():
				doc.set(key, value)
			doc.save(ignore_permissions=True)
			updated += 1
			continue
		doc = frappe.get_doc({"doctype": "PRAI FAQ", "title": title, **payload})
		doc.insert(ignore_permissions=True)
		created += 1
	return {"created": created, "updated": updated, "skipped": skipped}


def process_import_document(
	import_name: str,
	*,
	step: str = "preview",
) -> dict:
	doc = frappe.get_doc("PRAI FAQ Import", import_name)
	if not doc.import_file:
		frappe.throw(_("Attach a PDF or DOCX file first."), frappe.ValidationError)

	doc.file_name = Path(doc.import_file.split("/")[-1]).name
	text = extract_document_text(doc.import_file)
	doc.extracted_text = text[:20_000]
	doc.import_log = ""

	if step == "extract":
		doc.status = "Extracted"
		doc.save(ignore_permissions=True)
		frappe.db.commit()
		return {
			"success": True,
			"status": doc.status,
			"characters": len(text),
			"preview_chars": len(doc.extracted_text),
		}

	items = generate_faq_items_from_text(
		text,
		default_category=doc.default_category or "General",
		default_module_area=doc.default_module_area or "General",
		use_openai=cint(doc.use_openai),
	)
	doc.set("items", preview_import_lines(items))
	doc.status = "Preview Ready"
	doc.import_log = _(
		"Generated {0} FAQ proposal(s). Review rows, uncheck any to skip, then click "
		"'Save to PRAI FAQ (required for chat)' — until then PRAI Assistant cannot use these answers."
	).format(len(doc.items))
	doc.save(ignore_permissions=True)
	frappe.db.commit()
	return {
		"success": True,
		"status": doc.status,
		"generated": len(doc.items),
		"characters": len(text),
	}
